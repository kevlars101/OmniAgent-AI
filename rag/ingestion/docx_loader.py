import asyncio
from typing import Dict, Any
import docx
from pydantic import BaseModel
import logging
from datetime import datetime, timezone

logger = logging.getLogger(__name__)

class DocumentExtract(BaseModel):
    text: str
    metadata: Dict[str, Any]

class DocxLoader:
    """
    Production-grade DOCX loader using python-docx.
    """
    def __init__(self):
        pass

    async def load_async(self, file_path: str, user_id: str, document_id: str) -> DocumentExtract:
        logger.info(f"Loading DOCX: {file_path}")
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._load_sync, file_path, user_id, document_id)

    def _load_sync(self, file_path: str, user_id: str, document_id: str) -> DocumentExtract:
        try:
            doc = docx.Document(file_path)
            full_text = [para.text for para in doc.paragraphs]
            
            metadata = {
                "source": file_path,
                "document_id": str(document_id),
                "user_id": str(user_id),
                "file_type": "docx",
                "paragraph_count": len(doc.paragraphs),
                "created_at": datetime.now(timezone.utc).isoformat(),
                "ingestion_version": "1.0"
            }
            
            return DocumentExtract(text="\n".join(full_text), metadata=metadata)
        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise RuntimeError(f"Corrupted or invalid DOCX file: {file_path}") from e
